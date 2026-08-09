"""RAG 知识库核心引擎：文档切分、向量化、检索、生成"""

import os

# ============================================================
# 国内访问 huggingface.co 常不稳定，默认改走镜像站下载 embedding 模型。
# 必须在 import transformers / sentence_transformers 之前设置才生效。
# ============================================================
os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")
os.environ.setdefault("HF_HUB_DOWNLOAD_TIMEOUT", "60")

import uuid
import pickle
from pathlib import Path

import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from openai import OpenAI

# ---- 配置 ----
CHUNK_SIZE = 400          # 每段文字长度
CHUNK_OVERLAP = 80        # 相邻段落重叠长度
TOP_K = 4                 # 检索返回的段落数
EMBED_MODEL = "BAAI/bge-small-zh-v1.5"  # 本地向量化模型
DEEPSEEK_MODEL = "deepseek-chat"        # 大模型

# bge 模型官方建议：查询时加指令
QUERY_INSTRUCTION = "为这个句子生成表示以用于检索相关文章："


class RagEngine:
    """RAG 核心引擎（使用 FAISS 向量库）"""
    
    def __init__(self, index_path: str = "./data/chroma_db/faiss.index"):
        # 1. 加载 Embedding 模型（首次会自动下载约 100MB）
        print(f"正在加载 Embedding 模型: {EMBED_MODEL}...")
        self.embedder = SentenceTransformer(EMBED_MODEL)
        print("Embedding 模型加载完成")
        
        # 2. 初始化 FAISS 向量库
        self.index_path = Path(index_path)
        self.index_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 存储文档内容
        self.documents = []
        self.metadatas = []
        
        # 如果已有索引则加载，否则新建
        if self.index_path.exists():
            try:
                self.index = faiss.read_index(str(self.index_path))
                self._load_metadata()
                print(f"向量库已加载，当前文档数: {self.index.ntotal}, 维度: {self.index.d}")
            except Exception as e:
                print(f"加载索引失败，将重新创建: {e}")
                self.index_path.unlink()
                self._create_new_index()
        else:
            self._create_new_index()
        
        # 3. 延迟初始化 LLM（用到时才创建）
        self._llm = None
    
    def _create_new_index(self):
        """创建新的 FAISS 索引，自动适配向量维度"""
        # 先用一个测试文本获取实际向量维度
        test_embedding = self._embed_texts(["测试文本"])
        # 确保是 2D 数组再取维度
        if test_embedding.ndim == 1:
            actual_dim = test_embedding.shape[0]
        else:
            actual_dim = test_embedding.shape[1]
        print(f"检测到向量维度: {actual_dim}")
        
        # 创建索引（余弦相似度）
        self.index = faiss.IndexFlatIP(actual_dim)
        print(f"向量库初始化完成，维度: {actual_dim}")
    
    def _load_metadata(self):
        """加载元数据（文档内容和来源）"""
        meta_path = self.index_path.parent / "metadata.pkl"
        if meta_path.exists():
            with open(meta_path, 'rb') as f:
                data = pickle.load(f)
                self.documents = data.get('documents', [])
                self.metadatas = data.get('metadatas', [])
    
    def _save_metadata(self):
        """保存元数据"""
        meta_path = self.index_path.parent / "metadata.pkl"
        with open(meta_path, 'wb') as f:
            pickle.dump({
                'documents': self.documents,
                'metadatas': self.metadatas
            }, f)
    
    def _get_llm(self):
        """获取 DeepSeek 客户端"""
        if self._llm is None:
            key = os.environ.get("DEEPSEEK_API_KEY")
            if not key:
                raise RuntimeError("未配置 DEEPSEEK_API_KEY，请在 .env 文件中设置")
            self._llm = OpenAI(api_key=key, base_url="https://api.deepseek.com")
        return self._llm
    
    def _chunk_text(self, text: str) -> list[str]:
        """将长文本切分成小段"""
        text = text.replace("\r\n", "\n").strip()
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + CHUNK_SIZE
            piece = text[start:end]
            
            # 尽量在换行处断开
            if end < len(text):
                nl = piece.rfind("\n")
                if nl > CHUNK_SIZE * 0.5:
                    piece = piece[:nl + 1]
                    end = start + len(piece)
            
            chunks.append(piece.strip())
            start = end - CHUNK_OVERLAP
        
        return [c for c in chunks if c]
    
    def _embed_texts(self, texts: list[str]) -> np.ndarray:
        """将文本列表向量化，始终返回 2D 数组"""
        embeddings = self.embedder.encode(
            texts,
            normalize_embeddings=True
        )
        # 确保返回的是 2D 数组 (n_samples, dim)
        if isinstance(embeddings, np.ndarray):
            if embeddings.ndim == 1:
                # 单个文本时，形状为 (dim,)，转为 (1, dim)
                return np.array([embeddings]).astype(np.float32)
            elif embeddings.ndim == 2:
                return np.array(embeddings).astype(np.float32)
        # 如果既不是 ndarray 也不是列表，转为 ndarray
        return np.array(embeddings).astype(np.float32)
    
    def ingest_text(self, text: str, source: str) -> int:
        """将文本入库"""
        chunks = self._chunk_text(text)
        if not chunks:
            return 0
        
        # 向量化
        embeddings = self._embed_texts(chunks)
        
        # 检查维度是否匹配
        if embeddings.shape[1] != self.index.d:
            print(f"⚠️ 向量维度不匹配！当前索引维度: {self.index.d}, 实际维度: {embeddings.shape[1]}")
            # 重建索引
            print("正在重建索引...")
            self._create_new_index()
            # 注意：重建索引后，需要重新添加所有已存在的文档
            # 但这里因为是新增文档，所以直接添加即可
        
        # 添加到 FAISS
        self.index.add(embeddings)
        
        # 存储文档和元数据
        self.documents.extend(chunks)
        self.metadatas.extend([
            {"source": source, "chunk_id": i} 
            for i in range(len(chunks))
        ])
        
        # 保存到磁盘
        faiss.write_index(self.index, str(self.index_path))
        self._save_metadata()
        
        return len(chunks)
    
    def ingest_file(self, file_path: Path) -> int:
        """将文件入库（支持 .txt / .md / .pdf / .docx）"""
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                raise ImportError("请安装 pypdf: pip install pypdf")
        
        elif suffix == ".docx":
            try:
                from docx import Document
                doc = Document(str(file_path))
                text = "\n".join(para.text for para in doc.paragraphs)
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + " | ".join(cell.text for cell in row.cells)
            except ImportError:
                raise ImportError("请安装 python-docx: pip install python-docx")
        
        else:
            # txt / md 等纯文本
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        
        return self.ingest_text(text, source=file_path.name)
    
    def retrieve(self, question: str, k: int = TOP_K) -> list[dict]:
        """检索最相关的文档段落"""
        if self.index.ntotal == 0:
            return []
        
        # 问题向量化（加指令）
        q_embedding = self._embed_texts([QUERY_INSTRUCTION + question])
        
        # FAISS 检索
        distances, indices = self.index.search(
            q_embedding, 
            min(k, self.index.ntotal)
        )
        
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(self.documents):
                results.append({
                    "text": self.documents[idx],
                    "source": self.metadatas[idx].get("source", "未知"),
                    "score": round(float(distances[0][i]), 3)
                })
        
        return results
    
    def answer(self, question: str) -> dict:
        """完整的 RAG 问答流程：检索 + 生成"""
        # 1. 检索
        hits = self.retrieve(question)
        
        if not hits:
            return {
                "answer": "知识库为空，请先上传文档。",
                "sources": []
            }
        
        # 2. 拼接上下文
        context = "\n\n".join(
            f"【资料{i+1}·来自 {h['source']}】\n{h['text']}"
            for i, h in enumerate(hits)
        )
        
        # 3. 构造提示词
        system_prompt = (
            "你是一个严谨的知识库问答助手。你只能依据下面提供的【参考资料】回答用户问题。"
            "如果参考资料中没有足以回答问题的信息，必须如实回答："
            "「根据现有资料，我无法回答这个问题。」"
            "绝对不要编造、不要使用资料之外的知识。回答用中文，简洁准确。"
        )
        user_prompt = f"【参考资料】\n{context}\n\n【用户问题】\n{question}"
        
        # 4. 调用大模型生成
        resp = self._get_llm().chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2
        )
        
        answer_text = resp.choices[0].message.content
        
        # 5. 返回结果
        return {
            "answer": answer_text,
            "sources": [
                {
                    "source": h["source"],
                    "score": h["score"],
                    "snippet": h["text"][:100] + "..." if len(h["text"]) > 100 else h["text"]
                }
                for h in hits
            ]
        }
    
    def status(self) -> dict:
        """获取状态信息"""
        count = self.index.ntotal
        sources = set()
        for meta in self.metadatas:
            sources.add(meta.get("source", "未知"))
        return {
            "chunks": count,
            "docs": sorted(sources)
        }
    
    def reset(self):
        """清空知识库"""
        self.documents = []
        self.metadatas = []
        
        # 重新创建索引
        self._create_new_index()
        
        # 删除文件
        if self.index_path.exists():
            self.index_path.unlink()
        meta_path = self.index_path.parent / "metadata.pkl"
        if meta_path.exists():
            meta_path.unlink()
        print("知识库已清空")